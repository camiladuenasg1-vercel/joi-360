# -*- coding: utf-8 -*-
"""
Generador del Historial de Tareas y Commits JOI360.
Lee src/00_historial.md (generado por gen_tasks_doc.py a partir del tracker
de tareas + git log real) y arma un .docx real.

Uso:
    python build_docx.py

Para actualizar: correr gen_tasks_doc.py de nuevo (recalcula desde el git log
real) y luego este script. Subir VERSION en cada corte dominical.
"""
import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(BASE, "src")

VERSION = "1.0"
OUT_NAME = f"JOI360_Historial_Tareas_y_Commits_v{VERSION}.docx"
OUT_PATHS = [
    os.path.join(BASE, OUT_NAME),
    os.path.join(r"C:\Users\CamilaDueñas\OneDrive - RedPontis\Escritorio", OUT_NAME),
]

SECTION_FILES = ["00_historial.md"]

PRIMARY = RGBColor(0x00, 0x35, 0xB9)
DARK = RGBColor(0x0B, 0x1C, 0x30)
MUTED = RGBColor(0x44, 0x46, 0x55)

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)

for lvl, size, color in [(1, 18, PRIMARY), (2, 13, PRIMARY), (3, 11, DARK), (4, 10, DARK)]:
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(16 if lvl == 1 else 10)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.keep_with_next = True

title_style = doc.styles["Title"]
title_style.font.name = "Calibri"
title_style.font.size = Pt(28)
title_style.font.color.rgb = PRIMARY
title_style.font.bold = True

mono_style = doc.styles.add_style("CodeBlock", 1)
mono_style.font.name = "Consolas"
mono_style.font.size = Pt(9)
mono_style.font.color.rgb = MUTED

section = doc.sections[0]
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

footer_p = section.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run()
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
run._r.append(fld)
for r in footer_p.runs:
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")


def add_inline_runs(paragraph, text):
    for chunk in INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            r = paragraph.add_run(chunk[2:-2])
            r.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = paragraph.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xB0, 0x00, 0x50)
        elif chunk.startswith("*") and chunk.endswith("*") and not chunk.startswith("**"):
            r = paragraph.add_run(chunk[1:-1])
            r.italic = True
            r.font.color.rgb = MUTED
        else:
            paragraph.add_run(chunk)


def add_table(doc, rows):
    header, *body = rows
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "0035B9")
    for row in body:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if i >= len(cells):
                break
            cells[i].paragraphs[0].clear()
            p = cells[i].paragraphs[0]
            add_inline_runs(p, val)
            for r in p.runs:
                if r.font.size is None:
                    r.font.size = Pt(8.5)
                    r.font.name = "Consolas" if i == 1 and len(header) == 3 else r.font.name
    doc.add_paragraph()


def parse_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator_row(line):
    cells = parse_table_row(line)
    return all(re.fullmatch(r":?-{2,}:?", c) for c in cells)


first_h1_seen = False


def render_file(path):
    global first_h1_seen
    with open(path, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [parse_table_row(l) for l in table_lines if not is_separator_row(l)]
            if rows:
                add_table(doc, rows)
            continue

        if stripped == "---":
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1 and not first_h1_seen:
                p = doc.add_paragraph(text, style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_h1_seen = True
            elif level == 1:
                doc.add_page_break()
                doc.add_heading(text, level=1)
            else:
                doc.add_heading(text, level=level)
            i += 1
            continue

        if first_h1_seen and "Versión" in stripped and len(doc.paragraphs) < 4:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped)
            r.italic = True
            r.font.color.rgb = MUTED
            r.font.size = Pt(12)
            i += 1
            continue

        m_bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if m_bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, m_bullet.group(1))
            i += 1
            continue

        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() != "" and not re.match(r"^(#{1,4})\s", lines[i].strip()) \
                and not lines[i].strip().startswith("|") and not re.match(r"^[-*]\s+", lines[i].strip()):
            para_lines.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(para_lines))


for fname in SECTION_FILES:
    path = os.path.join(SRC, fname)
    if not os.path.exists(path):
        print(f"AVISO: falta {fname}")
        continue
    render_file(path)

for out_path in OUT_PATHS:
    doc.save(out_path)
    print("Guardado:", out_path)
