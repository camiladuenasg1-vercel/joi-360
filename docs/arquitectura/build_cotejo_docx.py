# -*- coding: utf-8 -*-
"""Genera el .docx del Cotejo Prototipo vs Proyecto Real desde su .md."""
import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(BASE, "cotejo_prototipo_vs_proyecto_real.md")
OUT_NAME = "JOI360_Cotejo_Prototipo_vs_ProyectoReal.docx"
OUT_PATHS = [
    os.path.join(BASE, OUT_NAME),
    os.path.join(r"C:\Users\CamilaDueñas\OneDrive - RedPontis\Escritorio", OUT_NAME),
]

PRIMARY = RGBColor(0x00, 0x35, 0xB9)
DARK = RGBColor(0x0B, 0x1C, 0x30)
MUTED = RGBColor(0x44, 0x46, 0x55)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)
for lvl, size, color in [(1, 19, PRIMARY), (2, 14, PRIMARY), (3, 12, DARK), (4, 11, DARK)]:
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = "Calibri"; st.font.size = Pt(size); st.font.color.rgb = color; st.font.bold = True
    st.paragraph_format.space_before = Pt(16 if lvl == 1 else 11)
    st.paragraph_format.space_after = Pt(6); st.paragraph_format.keep_with_next = True
ts = doc.styles["Title"]; ts.font.name = "Calibri"; ts.font.size = Pt(26); ts.font.color.rgb = PRIMARY; ts.font.bold = True
mono = doc.styles.add_style("CodeBlock", 1)
mono.font.name = "Consolas"; mono.font.size = Pt(8.5); mono.font.color.rgb = MUTED
mono.paragraph_format.left_indent = Cm(0.5); mono.paragraph_format.space_after = Pt(2)
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Cm(2.0); sec.top_margin = sec.bottom_margin = Cm(2)
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run(); fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); r._r.append(fld)
for rr in fp.runs: rr.font.size = Pt(8); rr.font.color.rgb = MUTED


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:fill"), hexc); tcPr.append(s)


INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def runs(p, text):
    for c in INLINE.split(text):
        if not c: continue
        if c.startswith("**") and c.endswith("**"):
            x = p.add_run(c[2:-2]); x.bold = True
        elif c.startswith("`") and c.endswith("`"):
            x = p.add_run(c[1:-1]); x.font.name = "Consolas"; x.font.size = Pt(9); x.font.color.rgb = RGBColor(0xB0, 0x00, 0x50)
        else:
            p.add_run(c)


def add_table(rows):
    head, *body = rows
    t = doc.add_table(rows=1, cols=len(head)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = True
    for i, h in enumerate(head):
        cel = t.rows[0].cells[i]; cel.paragraphs[0].clear()
        x = cel.paragraphs[0].add_run(h); x.bold = True; x.font.size = Pt(9); x.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cel, "0035B9")
    for row in body:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            if i >= len(cs): break
            cs[i].paragraphs[0].clear(); pp = cs[i].paragraphs[0]; runs(pp, v)
            for rr in pp.runs:
                if rr.font.size is None: rr.font.size = Pt(8.5)
    doc.add_paragraph()


def prow(line): return [c.strip() for c in line.strip().strip("|").split("|")]
def is_sep(line): return all(re.fullmatch(r":?-{2,}:?", c) for c in prow(line))


first_h1 = False
with open(SRC, encoding="utf-8") as f:
    lines = f.read().split("\n")
i, n = 0, len(lines)
while i < n:
    s = lines[i].strip()
    if s == "": i += 1; continue
    if s.startswith("|"):
        tl = []
        while i < n and lines[i].strip().startswith("|"): tl.append(lines[i].strip()); i += 1
        rr = [prow(l) for l in tl if not is_sep(l)]
        if rr: add_table(rr)
        continue
    if s.startswith("```"):
        i += 1; code = []
        while i < n and not lines[i].strip().startswith("```"): code.append(lines[i]); i += 1
        i += 1
        p = doc.add_paragraph(style="CodeBlock"); p.add_run("\n".join(code)); continue
    if s == "---": i += 1; continue
    m = re.match(r"^(#{1,4})\s+(.*)$", s)
    if m:
        lvl = len(m.group(1)); txt = m.group(2).strip()
        if lvl == 1 and not first_h1:
            p = doc.add_paragraph(txt, style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; first_h1 = True
        elif lvl == 1:
            doc.add_page_break(); doc.add_heading(txt, level=1)
        else:
            doc.add_heading(txt, level=lvl)
        i += 1; continue
    mb = re.match(r"^[-*]\s+(.*)$", s); mn = re.match(r"^\d+\.\s+(.*)$", s)
    if mb:
        runs(doc.add_paragraph(style="List Bullet"), mb.group(1)); i += 1; continue
    if mn:
        runs(doc.add_paragraph(style="List Number"), mn.group(1)); i += 1; continue
    para = [s]; i += 1
    while i < n and lines[i].strip() != "" and not re.match(r"^(#{1,4})\s", lines[i].strip()) \
            and not lines[i].strip().startswith("|") and not lines[i].strip().startswith("```") \
            and not re.match(r"^[-*]\s+", lines[i].strip()) and not re.match(r"^\d+\.\s+", lines[i].strip()):
        para.append(lines[i].strip()); i += 1
    runs(doc.add_paragraph(), " ".join(para))

for op in OUT_PATHS:
    doc.save(op); print("Guardado:", op)
